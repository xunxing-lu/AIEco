from __future__ import annotations
from contextlib import AsyncExitStack
from typing import Any, Dict, List
from dataclasses import dataclass
from dotenv import load_dotenv
from rich.markdown import Markdown
from rich.console import Console
from rich.live import Live
import asyncio
import os

from pydantic import BaseModel, Field
from pydantic_ai.providers.openai import OpenAIProvider
from pydantic_ai.models.openai import OpenAIModel
from pydantic_ai.mcp import MCPServerStdio
from pydantic_ai import Agent, RunContext

from docx import Document
from docx.shared import Inches
import docx
import re
import os

from pydantic_ai.models.gemini import GeminiModel
from pydantic_ai.providers.google_gla import GoogleGLAProvider

load_dotenv()

gavin_case = {
    'ctpt': '../data/physio/Gavin/gavinpt.txt',
    'solo': '../data/physio/Gavin/gavinsolo.txt',
    'template_file' : r"../data/physio/Progress_Note_Template.docx",
    'output_file' : r"../data/physio/Gavin/Progress_Note_Gavin_updated.docx"
}

margrate_case = {
    'ctpt': '../data/physio/Margrate/margaretpt.txt',
    'solo': '../data/physio/Margrate/margaretsolo.txt',
    'template_file' : r"../data/physio/Progress_Note_Template.docx",
    'output_file' : r"../data/physio/Margrate/Progress_Note_Margrate_updated.docx"
}

picked_case = gavin_case

def update_word_template(template_path, output_path, replacement_dict, image_replacements):
    doc = Document(template_path)
    pattern = r'\[(.*?)\]'

    def process_paragraphs(paragraphs):
        for paragraph in paragraphs:
            full_text = ''.join(run.text for run in paragraph.runs)

            if '[Scooter_Image]' in full_text:
                print("found image")


                # Clear all runs
                for run in paragraph.runs:
                    run.text = ''

                # Split around the image placeholder
                parts = full_text.split('[Scooter_Image]')
                
                # Rebuild the paragraph
                if parts[0]:
                    paragraph.add_run(parts[0])
                # Add the image
                paragraph.add_run().add_picture(image_replacements['Scooter_Image'], width=Inches(3))
                if len(parts) > 1:
                    paragraph.add_run(parts[1])
            else:
                for match in re.finditer(pattern, paragraph.text):
                    key = match.group(1)
                    if key in replacement_dict:
                        paragraph.text = paragraph.text.replace(f'[{key}]', replacement_dict[key])

    process_paragraphs(doc.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

    for section in doc.sections:
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer is not None:
                for p in footer.paragraphs:
                    print("Footer paragraph:", p.text)

                process_paragraphs(footer.paragraphs)
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            # for p in cell.paragraphs:
                            #     print("Cell paragraph:", p.text)
                            process_paragraphs(cell.paragraphs)

    # for section in doc.sections:
    #     footer = section.footer
    #     # for p in footer.paragraphs:
    #     #     print("Footer paragraph:", p.text)
        
    #     process_paragraphs(footer.paragraphs)

    #     # If your footer contains tables
    #     for table in footer.tables:
    #         for row in table.rows:
    #             for cell in row.cells:
    #                 # for p in cell.paragraphs:
    #                 #     print("Cell paragraph:", p.text)
    #                 process_paragraphs(cell.paragraphs)


    for section in doc.sections:
        header = section.header
        process_paragraphs(header.paragraphs)

        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    process_paragraphs(cell.paragraphs)



    doc.save(output_path)
    print(f"Document saved to {output_path}")


class Assessment(BaseModel):
    Assessment_Date: str = Field(description='date of assessment, use current date if not provided')
    Initial_Assessment_Info: str = Field(description='Brief client summary including age, gender, referral reason (e.g., PT under HCP), and primary condition to be managed')
    Subjective_Examination_Info: str = Field(description='Record of client consent for assessment, scribe use, and acknowledgment of COVID-19 precautions')
    Presenting_Issue_Info: str = Field(description='Client-reported functional difficulties and symptoms impacting daily activities, mobility, and participation in usual roles or routines. use paragraphs or multilines if needed for well organised format.')
    His_of_Presenting_Issue_Info: str = Field(description='Timeline and progression of symptoms or diagnosis, including onset, severity changes, and current status over time')
    Transers_Mobility_Info: str = Field( description='Client’s ability to perform transfers (e.g., sit-to-stand, bed, chair) and mobilise within the home and community, including use of aids, distance managed, rest breaks, and assistance required. use paragraphs or multilines if needed for well organised format.')
    Falls_Info: str = Field(description='History of falls and near misses, including frequency, timing, causes, injuries, hospitalisations, use of mobility aids, ability to recover from falls, and current fall risk. use paragraphs or multilines if needed for well organised format.')
    Pain_Info: str = Field(description='Details of pain including location, type, severity (e.g., VAS), duration, aggravating/easing factors, 24-hour pattern, functional impact, and any associated symptoms or current pain management. use paragraphs or multilines if needed for well organised format.')
    Social_Hx_Info: str = Field(description='Client’s social background including living situation, cultural background, family and community support, caregiving roles, social engagement, and any existing or required formal/informal assistance. use paragraphs or multilines or bullet points if needed for well organised format.')
    FormalA_AHP_Equi_Info: str = Field(description='Details of current or past involvement with allied health professionals (e.g., OT, PT), including referrals, engagement with therapy, home exercise programs, assistive equipment, home modifications, and client preferences or concerns regarding formal supports. use paragraphs or multilines or bullet points if needed for well organised format.')
    General_Health_Info: str = Field(description='Summary of relevant medical history, comorbidities, current health conditions, and functional impacts (e.g., incontinence, hearing, swallowing). Include any ongoing specialist care, medications, red flag clearance, and notable absences of conditions. use paragraphs or multilines if needed for well organised format.use paragraphs or multilines or bullet points if needed for well organised format.')
    Medications_Info: str = Field(description='Details of current medications (prescribed and over-the-counter), including any uncertainty or lack of recall. Note use of Webster packs, inhalers, or other aids, and include client consent to photograph medications if applicable. use paragraphs or multilines or bullet points if needed for well organised format.')
    Other_Concerns_Questions_Info: str = Field(description='Client-expressed questions, concerns, or beliefs about their condition, prognosis, or treatment (including exercise hesitancy, alternative therapies, or specialist care). Include any misinformation, requests for clarification, or advice provided within scope. use paragraphs or multilines if needed for well organised format.')
    Goals_Info: str = Field(description='Client-identified short- and long-term goals related to physical function, independence, pain reduction, and quality of life. Include specific functional targets (e.g., walking to toilet, reducing rest breaks), broader aspirations (e.g., avoiding residential care, attending social events), and any agreed timelines. use bullet points or paragraphs for clear organisation.')
    Objective_Exam_Info: str = Field(description='Objective clinical observations during assessment, including alertness and orientation (e.g., A+O to TPP), physical responses, fatigue, and tolerance to activity. Note any observable changes, effort levels, or need for rest. use bullet points or paragraphs if needed for clarity.')
    Observations_Info: str = Field(description='Clinician’s visual and physical observations during assessment, including posture, breathing patterns, muscle use, fatigue, appearance, use of assistive equipment, home modifications, and environmental barriers. Use bullet points or paragraphs to organize relevant findings clearly.')	
    Palpation_Info: str = Field(description='Findings from physical palpation during assessment, including areas of tenderness, muscle stiffness or tone changes, joint mobility, and any asymmetries or abnormalities noted. Use paragraphs or bullet points to clearly organize relevant tactile observations.')
    Range_of_Motion_Info: str = Field(description='Assessment of joint and spinal range of motion, noting any limitations, stiffness, effort required, pain, or asymmetries during active and passive movements. Include specific joint angles or movement degrees when available, and describe functional impact clearly. use multi columns to display both left side and right side if needed for clear organisation.')
    MMT_Info: str = Field(description='Manual Muscle Testing (MMT) findings including muscle strength grades (e.g., 3/5, 4/5), effort required, endurance, ability to maintain contractions, presence of fatigue post-testing, and any observed asymmetries or limitations. Include specific muscle groups or movements assessed. use multi columns to display both left side and right side if needed for clear organisation.')
    Balance_Info: str = Field(description='Assessment of client’s static and dynamic balance abilities, including timed balance tests (e.g., FTEO, FTEC, FAEO, FAEC) with or without gait aids, ability to maintain standing balance safely, duration held, and use of assistive devices. Note any limitations or safety concerns observed during testing.use bullet points for clear organisation.')
    Transfers_Info: str = Field( description='Client’s ability to perform transfers (e.g., sit-to-stand, bed, chair), including level of physical assistance required, use of upper limbs or momentum, control during movements (e.g., eccentric control), presence of symptoms such as shortness of breath or fatigue post-transfer, number of attempts needed, and rest breaks required.')
    Gait_Mobility_Info: str = Field(description='Client’s gait and mobility characteristics including use of aids (e.g., 4-wheeled walker), gait patterns (step length, foot clearance, cadence, base of support), posture, gait disturbances (freezing, shuffling, step-to/step-through), ability to maintain direction and assistance needed, distance managed, fatigue or shortness of breath during ambulation, and required rest breaks.')
    STS_Info: str = Field(description = 'Client’s sit-to-stand (STS) performance including repetitions and time, use of upper limbs for assistance, chair type, effort level, breathlessness or fatigue, and rest break requirements.')
    Treatment_Info: str = Field(description='Summary of treatment provided or planned, including assessment status, home exercise program (HEP), education, goal discussion, and next steps.')
    Assessment_Info: str = Field(description='Brief summary of client’s overall condition, living situation, diagnosis, key functional limitations, and primary goals identified at assessment.use bullet points for clear organisation.')
    Plan_Info: str = Field(description='Planned interventions and follow-up actions including review schedule, referrals, exercise programs, equipment or document requests, communication with care team, and treatment priorities. Use bullet points for clear organisation.')
    TUGT_Info: str = Field(description='Timed Up and Go Test (TUGT) findings, including total time taken, use of mobility aids (e.g., 4WW), movement quality, presence of symptoms such as shortness of breath or fatigue, rest break requirements, and overall safety or effort observed during the test. Use bullet points or structured sentences for clarity.')
    MWT10_Info: str = Field(description='10-Meter Walk Test (10MWT) findings, including walking speed (m/s), use of mobility aids (e.g., 4WW), gait quality, symptoms during or post-test (e.g., breathlessness, fatigue), need for rest breaks, and observed safety or effort levels. Use structured sentences or bullet points for clarity.')
    Education_Info: str = Field(description='Summary of client education provided during assessment, including explanations of diagnosis, condition management, role of physiotherapy, and the nature and purpose of prescribed exercises. Note client’s level of understanding, emotional response (e.g., reassurance, anxiety), and engagement or receptiveness to information shared. Use paragraphs or structured points for clarity.')
    Diaphragmatic_Breathing_Info: str = Field(description='Assessment of client’s ability to perform diaphragmatic breathing, including use of physical cues (e.g., relaxing shoulders, tummy rise), level of prompting required, ability to replicate technique independently, and overall effectiveness. Note consistency, responsiveness to cues, and observed breathing patterns. Use structured sentences or bullet points for clarity.')
    PLB_Info: str = Field(description='Assessment of client’s ability to perform pursed-lip breathing (PLB), including use of physical cues (e.g., “kissing lips” technique), level of prompting required, and ability to reproduce the technique independently. Note consistency of demonstration, effectiveness, and client responsiveness to cueing. Use structured sentences or bullet points for clarity.')
    Physio_Info: str = Field(description='Physiotherapist\'s name')
    Physio_title_Info: str = Field(description='Physiotherapist\'s title')

# ========== Helper function to get model configuration ==========
def get_model():
    llm = 'o3'
    print(llm)
    base_url = 'https://api.openai.com/v1'
    print(base_url)
    api_key = os.getenv("OPENAI_API_KEY")
    print(api_key)
    return OpenAIModel(llm, provider=OpenAIProvider(base_url=base_url, api_key=api_key))

def get_gemini_model():
    api_key = os.getenv("GEMINI_API_KEY")

    return GeminiModel(
        'gemini-2.0-flash', 
        provider=GoogleGLAProvider(api_key=api_key)
    )

def read_word(file_path):
    """
    read text from word
    """
    try:
        # Load the Word document
        doc = docx.Document(file_path)
        
        # Extract all text from the document
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        
        text = '\n'.join(full_text)
        
        return text
    
    except Exception as e:
        return f"Error processing file: {str(e)}"


conversation1 = ''
with open(picked_case['ctpt'], 'r', encoding='utf-8') as file:
    conversation1 = file.read()

conversation2 = ''
with open(picked_case['solo'], 'r', encoding='utf-8') as file:
    file_content = file.read()

sample_progress_note_1 = read_word('../data/physio/Gavin/Gavin Progress Note.docx')
sample_progress_note_2 = read_word('../data/physio/Margrate/Margaret Demo Progress Note.docx')

primary_agent = Agent(
    get_model(),
    # get_gemini_model(),
    system_prompt=f"""
    You are a Senior Physiotherapist who is very experienced in Physio Assessment with patient and writing Progress Note.
    You will have conversation with patient, then have solo conversation to remember key content.
    
    Here are some sample Progress Notes for you to learn about:
        1, Sample Progress Note for study: {sample_progress_note_1} .
        2, Sample Progress Note for study: {sample_progress_note_2} .

    You need to use terminology in the progress note to be professional and simplify description wherever it is possible, here are some terminoloy for you to check:
    1, FAEO nil gait aide 7s which means the patient stands with their feet apart and their eyes open, did not require any walking aid, maintain this position safely for 7 seconds.
    2, FAEC nil gait aide 2s< which means the patient stands with their feet apart and their eyes open, did not require any walking aid, maintain this position safely for less than 2 seconds.
    3, FTEO 20s which means the client could stand with feet together, eyes open, for 20 seconds.
    4, FTEC 5s< which means the client could stand for less than 5 seconds with feet together and eyes closed.
    5, UL grossly 4/5 which means 
    "UL = Upper Limbs
    Grade	Description
    5/5	Normal strength (full resistance)
    4/5	Good strength (can move against moderate resistance)
    3/5	Fair (can move against gravity only)
    2/5	Poor (can move with gravity eliminated)
    1/5	Trace movement only
    0/5	No movement detected."
    6, TF which means the task/function.
    7, 5x STS 26 sec which means Five Times Sit-to-Stand Test, he client took 26 seconds to complete the 5 repetitions.
    8, 4WW which means 4-Wheeled Walker.
    9, WC which means Wheelchair.
    10, R/V 1/52 which means Review in 1 week.
    11, A+O which means Alert and Oriented.
    12, TPP which means Time, Place, Person. 
    13, I/M which means Intermittent.
    14, Dx which means Diagnosis.
    15, Hx which means History.
    16, a/a which means as above.
    17, Formal Dx ~6/12 ago which means formal diagnosis was made about 6 months prior to the note.
    18, Cx which means Cervical, referring to the cervical spine or neck region.
    19, Tx which means Thoracic, referring to the thoracic spine or mid-back region.
    20, Lx which means Lumbar, referring to the lumbar spine or lower back region.
    21, ROM which means Range of Motion, which is a measure of the movement around a specific joint or body part.
    22, A = 80/80 which means the individual can actively reach 80 degrees of movement, which is considered normal in this context since it matches the expected full range of 80 degrees.
    23, P = 80/80 which means passively, the joint can also reach the full 80-degree range.
    24, A = ½ which means whatever quantity depends on A is halved.
    25, Modified 30s : 4x which means a modified version of a sit-to-stand exercise where the patient performs the exercise four times within a 30-second period.
    26, 0.47 m/s which means a pace of 0.47 meters per second.
    27, Ax which means Assessment.
    28, BBS which means Berg Balance Scale.
    29, PMS which means Physical Mobility Scale.
    30, Cont. which means Continue.
    31, PMHx which means Past Medical History.
    32, SC which means Specialist Consultant or Service Coordinator, depending on the context.
    33, PMD which means Personal Mobility Device.
    34, ERC which means Equipment Resource Center.
    35, PMS which means Performance Measurement System or Patient Management System, depending on the context.
    36, CV which means Cardiovascular.
    37, OT which means Occupational Therapy.
    38, RV which means Review.
    39, PT which means Physiotherapy (or Physical Therapy).
    40, COPD which means Chronic Obstructive Pulmonary Disease, a chronic inflammatory lung disease that causes obstructed airflow from the lungs. It's characterized by long-term breathing problems and poor airflow.
    41, ADLs which means Activities of Daily Living, which refer to the basic tasks necessary for everyday functioning, such as eating, bathing, dressing, toileting, and moving around.
    42, SOB which means Shortness of Breath, a common symptom associated with many conditions, including COPD. It refers to the feeling of not being able to breathe well or having difficulty breathing.
    43, OM which means Outcome Measures, which are tools used to assess the patient's current status, treatment effectiveness, and progress over time. They are often standardized tests or questionnaires.

    by looking at the conversations, you need to extract the information and fill in the following fields:

    1, Assessment_Date: 'date of assessment, use current date if not provided'
    2, Initial_Assessment_Info: 'Brief client summary including age, gender, referral reason (e.g., PT under HCP), and primary condition to be managed'
    3, Subjective_Examination_Info: 'Record of client consent for assessment, scribe use, and acknowledgment of COVID-19 precautions'
    4, Presenting_Issue_Info: 'Client-reported functional difficulties and symptoms impacting daily activities, mobility, and participation in usual roles or routines. use paragraphs or multilines if needed for well organised format.'
    5, His_of_Presenting_Issue_Info: 'Timeline and progression of symptoms or diagnosis, including onset, severity changes, and current status over time'
    6, Transers_Mobility_Info: 'Client\'s ability to perform transfers (e.g., sit-to-stand, bed, chair) and mobilise within the home and community, including use of aids, distance managed, rest breaks, and assistance required. use paragraphs or multilines if needed for well organised format.'
    7, Falls_Info: 'History of falls and near misses, including frequency, timing, causes, injuries, hospitalisations, use of mobility aids, ability to recover from falls, and current fall risk. use paragraphs or multilines if needed for well organised format.'
    8, Pain_Info: 'Details of pain including location, type, severity (e.g., VAS), duration, aggravating/easing factors, 24-hour pattern, functional impact, and any associated symptoms or current pain management. use paragraphs or multilines if needed for well organised format.'
    9, Social_Hx_Info: 'Client\'s social background including living situation, cultural background, family and community support, caregiving roles, social engagement, and any existing or required formal/informal assistance. use paragraphs or multilines or bullet points if needed for well organised format.'
    10, FormalA_AHP_Equi_Info: 'Details of current or past involvement with allied health professionals (e.g., OT, PT), including referrals, engagement with therapy, home exercise programs, assistive equipment, home modifications, and client preferences or concerns regarding formal supports. use paragraphs or multilines or bullet points if needed for well organised format.'
    11, General_Health_Info: 'Summary of relevant medical history, comorbidities, current health conditions, and functional impacts (e.g., incontinence, hearing, swallowing). Include any ongoing specialist care, medications, red flag clearance, and notable absences of conditions. use paragraphs or multilines if needed for well organised format.use paragraphs or multilines or bullet points if needed for well organised format.'
    12, Medications_Info: 'Details of current medications (prescribed and over-the-counter), including any uncertainty or lack of recall. Note use of Webster packs, inhalers, or other aids, and include client consent to photograph medications if applicable. use paragraphs or multilines or bullet points if needed for well organised format.'
    13, Other_Concerns_Questions_Info: 'Client-expressed questions, concerns, or beliefs about their condition, prognosis, or treatment (including exercise hesitancy, alternative therapies, or specialist care). Include any misinformation, requests for clarification, or advice provided within scope. use paragraphs or multilines if needed for well organised format.'
    14, Goals_Info: 'Client-identified short- and long-term goals related to physical function, independence, pain reduction, and quality of life. Include specific functional targets (e.g., walking to toilet, reducing rest breaks), broader aspirations (e.g., avoiding residential care, attending social events), and any agreed timelines. use bullet points or paragraphs for clear organisation.'
    15, Objective_Exam_Info: 'Objective clinical observations during assessment, including alertness and orientation (e.g., A+O to TPP), physical responses, fatigue, and tolerance to activity. Note any observable changes, effort levels, or need for rest. use bullet points or paragraphs if needed for clarity.'
    16, Observations_Info: 'Clinician\'s visual and physical observations during assessment, including posture, breathing patterns, muscle use, fatigue, appearance, use of assistive equipment, home modifications, and environmental barriers. Use bullet points or paragraphs to organize relevant findings clearly.'	
    17, Palpation_Info: 'Findings from physical palpation during assessment, including areas of tenderness, muscle stiffness or tone changes, joint mobility, and any asymmetries or abnormalities noted. Use paragraphs or bullet points to clearly organize relevant tactile observations.'
    18, Range_of_Motion_Info: 'Assessment of joint and spinal range of motion, noting any limitations, stiffness, effort required, pain, or asymmetries during active and passive movements. Include specific joint angles or movement degrees when available, and describe functional impact clearly. use multi columns to display both left side and right side if needed for clear organisation.'
    19, MMT_Info: 'Manual Muscle Testing (MMT) findings including muscle strength grades (e.g., 3/5, 4/5), effort required, endurance, ability to maintain contractions, presence of fatigue post-testing, and any observed asymmetries or limitations. Include specific muscle groups or movements assessed. use multi columns to display both left side and right side if needed for clear organisation.'
    20, Balance_Info: 'Assessment of client\'s static and dynamic balance abilities, including timed balance tests (e.g., FTEO, FTEC, FAEO, FAEC) with or without gait aids, ability to maintain standing balance safely, duration held, and use of assistive devices. Note any limitations or safety concerns observed during testing.use bullet points for clear organisation.'
    21, Transfers_Info: 'Client\'s ability to perform transfers (e.g., sit-to-stand, bed, chair), including level of physical assistance required, use of upper limbs or momentum, control during movements (e.g., eccentric control), presence of symptoms such as shortness of breath or fatigue post-transfer, number of attempts needed, and rest breaks required.'
    22, Gait_Mobility_Info: 'Client\'s gait and mobility characteristics including use of aids (e.g., 4-wheeled walker), gait patterns (step length, foot clearance, cadence, base of support), posture, gait disturbances (freezing, shuffling, step-to/step-through), ability to maintain direction and assistance needed, distance managed, fatigue or shortness of breath during ambulation, and required rest breaks.'
    23, STS_Info: 'Client\'s sit-to-stand (STS) performance including repetitions and time, use of upper limbs for assistance, chair type, effort level, breathlessness or fatigue, and rest break requirements.'
    24, Treatment_Info: 'Summary of treatment provided or planned, including assessment status, home exercise program (HEP), education, goal discussion, and next steps.'
    25, Assessment_Info: 'Brief summary of client\'s overall condition, living situation, diagnosis, key functional limitations, and primary goals identified at assessment.use bullet points for clear organisation.'
    26, Plan_Info: 'Planned interventions and follow-up actions including review schedule, referrals, exercise programs, equipment or document requests, communication with care team, and treatment priorities. Use bullet points for clear organisation.'
    27, TUGT_Info: 'Timed Up and Go Test (TUGT) findings, including total time taken, use of mobility aids (e.g., 4WW), movement quality, presence of symptoms such as shortness of breath or fatigue, rest break requirements, and overall safety or effort observed during the test. Use bullet points or structured sentences for clarity.'
    28, MWT10_Info: '10-Meter Walk Test (10MWT) findings, including walking speed (m/s), use of mobility aids (e.g., 4WW), gait quality, symptoms during or post-test (e.g., breathlessness, fatigue), need for rest breaks, and observed safety or effort levels. Use structured sentences or bullet points for clarity.'
    29, Education_Info: 'Summary of client education provided during assessment, including explanations of diagnosis, condition management, role of physiotherapy, and the nature and purpose of prescribed exercises. Note client\'s level of understanding, emotional response (e.g., reassurance, anxiety), and engagement or receptiveness to information shared. Use paragraphs or structured points for clarity.'
    30, Diaphragmatic_Breathing_Info: 'Assessment of client\'s ability to perform diaphragmatic breathing, including use of physical cues (e.g., relaxing shoulders, tummy rise), level of prompting required, ability to replicate technique independently, and overall effectiveness. Note consistency, responsiveness to cues, and observed breathing patterns. Use structured sentences or bullet points for clarity.'
    31, PLB_Info: 'Assessment of client\'s ability to perform pursed-lip breathing (PLB), including use of physical cues (e.g., “kissing lips” technique), level of prompting required, and ability to reproduce the technique independently. Note consistency of demonstration, effectiveness, and client responsiveness to cueing. Use structured sentences or bullet points for clarity.'
    32, Physio_Info: 'Physiotherapist\'s name'
    33, Physio_title_Info: 'Physiotherapist\'s title'
    """,
    result_type=Assessment,
)

content = f"""
    Here are two conversations:
        Conversation 1, which is between a physiotherapist and a patient: {conversation1} .
        Conversation 2, which is physiotherapist solo: {conversation2} .
"""
# print(file_content)

result = primary_agent.run_sync(content)
assess_result = result.data

print(assess_result.Initial_Assessment_Info)

# Dictionary for text replacements
replacements = {
    'Assessment_Date': assess_result.Assessment_Date,
    'Initial_Assessment_Info': assess_result.Initial_Assessment_Info,
    'Subjective_Examination_Info': assess_result.Subjective_Examination_Info,
    'Presenting_Issue_Info': assess_result.Presenting_Issue_Info,
    'His_of_Presenting_Issue_Info': assess_result.His_of_Presenting_Issue_Info,
    'Transers_Mobility_Info': assess_result.Transers_Mobility_Info,
    'Falls_Info': assess_result.Falls_Info,
    'Pain_Info': assess_result.Pain_Info,
    'Social_Hx_Info': assess_result.Social_Hx_Info,
    'FormalA_AHP_Equi_Info': assess_result.FormalA_AHP_Equi_Info,
    'General_Health_Info': assess_result.General_Health_Info,
    'Medications_Info': assess_result.Medications_Info,
    'Other_Concerns_Questions_Info': assess_result.Other_Concerns_Questions_Info,
    'Goals_Info': assess_result.Goals_Info, 
    'Objective_Exam_Info': assess_result.Objective_Exam_Info,
    'Observations_Info': assess_result.Observations_Info,   
    'Palpation_Info': assess_result.Palpation_Info,
    'Range_of_Motion_Info': assess_result.Range_of_Motion_Info,
    'MMT_Info': assess_result.MMT_Info,
    'Balance_Info': assess_result.Balance_Info,
    'Transfers_Info': assess_result.Transfers_Info,
    'Gait_Mobility_Info': assess_result.Gait_Mobility_Info,
    'STS_Info': assess_result.STS_Info,
    'Treatment_Info': assess_result.Treatment_Info,
    'Assessment_Info': assess_result.Assessment_Info,
    'Plan_Info': assess_result.Plan_Info,
    'TUGT_Info': assess_result.TUGT_Info,
    'MWT10_Info': assess_result.MWT10_Info,
    'Education_Info': assess_result.Education_Info,
    'Diaphragmatic_Breathing_Info': assess_result.Diaphragmatic_Breathing_Info,
    'PLB_Info': assess_result.PLB_Info,
    'Physio_Info': assess_result.Physio_Info,
    'Physio_title_Info': assess_result.Physio_title_Info
}

# Provide image path for [Scooter_Image]
image_replacements = {
    'Signature_Image': '../img/scooter.png'  # Replace with your actual image path
}

# File paths
input_path = picked_case['template_file']
output_path = picked_case['output_file']

update_word_template(input_path, output_path, replacements, image_replacements)
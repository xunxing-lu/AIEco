from __future__ import annotations
from contextlib import AsyncExitStack
from typing import Any, Dict, List
from dataclasses import dataclass
from dotenv import load_dotenv
from rich.markdown import Markdown
from rich.console import Console
from rich.live import Live
import asyncio
import os

from pydantic import BaseModel, Field
from pydantic_ai.providers.openai import OpenAIProvider
from pydantic_ai.models.openai import OpenAIModel
from pydantic_ai.mcp import MCPServerStdio
from pydantic_ai import Agent, RunContext

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Optional, Dict, Any
import docx
import re
import os

from pydantic_ai.models.gemini import GeminiModel
from pydantic_ai.providers.google_gla import GoogleGLAProvider

import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

load_dotenv()

gavin_case = {
    'ctpt': '../data/physio/Gavin/gavinpt.txt',
    'solo': '../data/physio/Gavin/gavinsolo.txt',
    'template_file' : r"../data/physio/Progress_Note_Template.docx",
    'output_file' : r"../data/physio/Gavin/Progress_Note_Gavin_updated.docx"
}

margrate_case = {
    'ctpt': '../data/physio/Margrate/margaretpt.txt',
    'solo': '../data/physio/Margrate/margaretsolo.txt',
    'template_file' : r"../data/physio/Progress_Note_Template.docx",
    'output_file' : r"../data/physio/Margrate/Progress_Note_Margrate_updated.docx"
}

test_case = {
    'ctpt': '../data/physio/Test/testpt.txt',
    'solo': '../data/physio/Test/testsolo.txt',
    'template_file' : r"../data/physio/Progress_Note_Template.docx",
    'output_file' : r"../data/physio/Test/Progress_Note_Test_updated.docx"
}

def get_g_model():
    llm = 'o3'
    # print(llm)
    base_url = 'https://api.openai.com/v1'
    # print(base_url)
    api_key = os.getenv("OPENAI_API_KEY")
    # print(api_key)
    return OpenAIModel(llm, provider=OpenAIProvider(base_url=base_url, api_key=api_key))

def get_o_model():
    llm = 'google/gemini-2.5-pro'
    # logger.info(f"Using model: {llm}")
    base_url = 'https://openrouter.ai/api/v1'
    api_key = os.getenv("OPEN_ROUTER_API_KEY")
    return OpenAIModel(llm, provider=OpenAIProvider(base_url=base_url, api_key=api_key))



picked_case = test_case
selected_model = get_o_model()


conversation1 = ''
with open(picked_case['ctpt'], 'r', encoding='utf-8') as file:
    conversation1 = file.read()

conversation2 = ''
with open(picked_case['solo'], 'r', encoding='utf-8') as file:
    file_content = file.read()

template_content=''
# Method 1: Using with statement (recommended)
with open('../data/physio/progress note template.txt', 'r') as file:
    template_content = file.read()

class ProgressNoteData(BaseModel):
    title: Optional[str] = None
    subtitle: Optional[str] = None
    paragraphs: List[str] = []
    bullet_points: Optional[List[str]] = None
    numbered_list: Optional[List[str]] = None
    table: Optional[Dict[str, Any]] = None

primary_agent = Agent(
    selected_model,
    result_type=ProgressNoteData,
    # get_gemini_model(),
    system_prompt=f"""
    You are a Senior Physiotherapist who is very experienced in Physio Assessment with patient and writing Progress Note.
    You will have conversation with patient, then have solo conversation to remember key content.

    You need to read the information from the conversion context provided and fill in the template provided later on and return the result in a word format data:
    
    1, fill in content within [content] by reading the explation within [content]
    
        such as:
        
        [Clinic Letterhead]

        [Clinic Address Line 1]
        [Clinic Address Line 2]
        [Contact Number]
        [Fax Number]

        will be:

        ABC Physiotherapy Clinic
        123 Main Street
        City, State, ZIP
        (123) 456-7890
        (123) 456-7891

    2, write a new paragrah by reading the explation within (explanation),

        such as:

        [Date of Note] (use format: DD Month YYYY)

        [Introduction] (Begin with a brief description of the patient, including their age, marital status, and living situation. This section must be written in full sentences as a cohesive paragraph. Do not use bullet points or lists.)

        will be:
        Date of Note: 15 October 2023

        Introduction
        Mr. John Doe is a 65-year-old, etc.  

    Note: You need to use professional terminology words in the progress note and simplify description wherever it is possible, here are some terminoloy for you to reference:
    1, FAEO nil gait aide 7s which means the patient stands with their feet apart and their eyes open, did not require any walking aid, maintain this position safely for 7 seconds.
    2, FAEC nil gait aide 2s< which means the patient stands with their feet apart and their eyes open, did not require any walking aid, maintain this position safely for less than 2 seconds.
    3, FTEO 20s which means the client could stand with feet together, eyes open, for 20 seconds.
    4, FTEC 5s< which means the client could stand for less than 5 seconds with feet together and eyes closed.
    5, UL grossly 4/5 which means 
    "UL = Upper Limbs
    Grade	Description
    5/5	Normal strength (full resistance)
    4/5	Good strength (can move against moderate resistance)
    3/5	Fair (can move against gravity only)
    2/5	Poor (can move with gravity eliminated)
    1/5	Trace movement only
    0/5	No movement detected."
    6, TF which means the task/function.
    7, 5x STS 26 sec which means Five Times Sit-to-Stand Test, he client took 26 seconds to complete the 5 repetitions.
    8, 4WW which means 4-Wheeled Walker.
    9, WC which means Wheelchair.
    10, R/V 1/52 which means Review in 1 week.
    11, A+O which means Alert and Oriented.
    12, TPP which means Time, Place, Person. 
    13, I/M which means Intermittent.
    14, Dx which means Diagnosis.
    15, Hx which means History.
    16, a/a which means as above.
    17, Formal Dx ~6/12 ago which means formal diagnosis was made about 6 months prior to the note.
    18, Cx which means Cervical, referring to the cervical spine or neck region.
    19, Tx which means Thoracic, referring to the thoracic spine or mid-back region.
    20, Lx which means Lumbar, referring to the lumbar spine or lower back region.
    21, ROM which means Range of Motion, which is a measure of the movement around a specific joint or body part.
    22, A = 80/80 which means the individual can actively reach 80 degrees of movement, which is considered normal in this context since it matches the expected full range of 80 degrees.
    23, P = 80/80 which means passively, the joint can also reach the full 80-degree range.
    24, A = ½ which means whatever quantity depends on A is halved.
    25, Modified 30s : 4x which means a modified version of a sit-to-stand exercise where the patient performs the exercise four times within a 30-second period.
    26, 0.47 m/s which means a pace of 0.47 meters per second.
    27, Ax which means Assessment.
    28, BBS which means Berg Balance Scale.
    29, PMS which means Physical Mobility Scale.
    30, Cont. which means Continue.
    31, PMHx which means Past Medical History.
    32, SC which means Specialist Consultant or Service Coordinator, depending on the context.
    33, PMD which means Personal Mobility Device.
    34, ERC which means Equipment Resource Center.
    35, PMS which means Performance Measurement System or Patient Management System, depending on the context.
    36, CV which means Cardiovascular.
    37, OT which means Occupational Therapy.
    38, RV which means Review.
    39, PT which means Physiotherapy (or Physical Therapy).
    40, COPD which means Chronic Obstructive Pulmonary Disease, a chronic inflammatory lung disease that causes obstructed airflow from the lungs. It's characterized by long-term breathing problems and poor airflow.
    41, ADLs which means Activities of Daily Living, which refer to the basic tasks necessary for everyday functioning, such as eating, bathing, dressing, toileting, and moving around.
    42, SOB which means Shortness of Breath, a common symptom associated with many conditions, including COPD. It refers to the feeling of not being able to breathe well or having difficulty breathing.
    43, OM which means Outcome Measures, which are tools used to assess the patient's current status, treatment effectiveness, and progress over time. They are often standardized tests or questionnaires.

    IMPORTANT: Return your response as a structured format with:
    - title: The main title of the progress note
    - paragraphs: List of paragraph texts for the progress note
    - Any other structured data as needed
    
    Format the progress note content as clear, professional paragraphs.
    """
)

content = f"""
    Here are two conversations:
        Conversation 1, which is between a physiotherapist and a patient: {conversation1} .
        Conversation 2, which is physiotherapist solo: {conversation2} .
    
    read the template provided below:
    
    Template:
    {template_content}
"""
# print(file_content)

def create_docx_file(filename, data):
    """
    Creates a new DOCX file with given data (overwrites if exists)
    
    Args:
        filename (str): Name of the DOCX file to create
        data (dict): Dictionary containing the document data
    """
    # Create a new document
    doc = Document()
    
    # Add title if provided
    if data.get('title'):
        title = doc.add_heading(data['title'], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add subtitle if provided
    if data.get('subtitle'):
        subtitle = doc.add_heading(data['subtitle'], level=1)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add paragraphs
    if data.get('paragraphs'):
        for paragraph_text in data['paragraphs']:
            if paragraph_text:  # Only add non-empty paragraphs
                doc.add_paragraph(paragraph_text)
    
    # Add bullet points (check if not None and not empty)
    if data.get('bullet_points'):
        for point in data['bullet_points']:
            if point:  # Only add non-empty points
                doc.add_paragraph(point, style='List Bullet')
    
    # Add numbered list (check if not None and not empty)
    if data.get('numbered_list'):
        for item in data['numbered_list']:
            if item:  # Only add non-empty items
                doc.add_paragraph(item, style='List Number')
    
    # Add table if provided
    if data.get('table'):
        table_data = data['table']
        if table_data.get('headers') and table_data.get('rows'):
            table = doc.add_table(rows=1, cols=len(table_data['headers']))
            table.style = 'Table Grid'
            
            # Add headers
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(table_data['headers']):
                hdr_cells[i].text = str(header)
            
            # Add data rows
            for row_data in table_data['rows']:
                row_cells = table.add_row().cells
                for i, cell_data in enumerate(row_data):
                    row_cells[i].text = str(cell_data)
    
    # Save the document (overwrites if exists)
    doc.save(filename)
    print(f"Document '{filename}' created successfully!")

result = primary_agent.run_sync(content)
assess_result = result.data
print(f"Progress Note Data: {assess_result}")

# File paths
output_path = picked_case['output_file']

create_docx_file(output_path, assess_result.model_dump())
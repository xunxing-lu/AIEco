from __future__ import annotations
from contextlib import AsyncExitStack
from typing import Any, Dict, List
from dataclasses import dataclass
from dotenv import load_dotenv
from rich.markdown import Markdown
from rich.console import Console
from rich.live import Live
import asyncio
import os

from pydantic import BaseModel, Field
from pydantic_ai.providers.openai import OpenAIProvider
from pydantic_ai.models.openai import OpenAIModel
from pydantic_ai.mcp import MCPServerStdio
from pydantic_ai import Agent, RunContext

from docx import Document
from docx.shared import Inches
import docx
import re
import os

load_dotenv()


gavin_case = {
    'ctpt': '../data/physio/Gavin/gavinpt.txt',
    'solo': '../data/physio/Gavin/gavinsolo.txt',
    'progress_note' : r"../data/physio/Gavin/Progress_Note_Gavin_updated.docx",
    'output_file' : r"../data/physio/Gavin/Email_Gavin.docx"
}

margrate_case = {
    'ctpt': '../data/physio/Margrate/margaretpt.txt',
    'solo': '../data/physio/Margrate/margaretsolo.txt',
    'progress_note' : r"../data/physio/Margrate/Progress_Note_Margrate_updated.docx",
    'output_file' : r"../data/physio/Margrate/Email_Margrate.docx"
}

picked_case = gavin_case

def write_text_to_word(text, filename="document.docx", title=None):
    """
    Write text to a Word document.
    
    Args:
        text (str): The text content to write to the document
        filename (str): Name of the output file (default: "document.docx")
        title (str, optional): Optional title for the document
    
    Returns:
        str: Path to the created document
    """
    # Create a new Document
    doc = Document()
    
    # Add title if provided
    if title:
        title_paragraph = doc.add_heading(title, 0)
    
    # Add the main text content
    # Split text by paragraphs (double newlines) for better formatting
    paragraphs = text.split('\n\n')
    
    for paragraph_text in paragraphs:
        if paragraph_text.strip():  # Skip empty paragraphs
            doc.add_paragraph(paragraph_text.strip())
    
    # Save the document
    doc.save(filename)
    
    return filename

# ========== Helper function to get model configuration ==========
def get_model():
    llm = 'o3'
    print(llm)
    base_url = 'https://api.openai.com/v1'
    print(base_url)
    api_key = os.getenv("OPENAI_API_KEY")
    print(api_key)
    return OpenAIModel(llm, provider=OpenAIProvider(base_url=base_url, api_key=api_key))

def read_word(file_path):
    """
    read text from word
    """
    try:
        # Load the Word document
        doc = docx.Document(file_path)
        
        # Extract all text from the document
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        
        text = '\n'.join(full_text)
        
        return text
    
    except Exception as e:
        return f"Error processing file: {str(e)}"
    
conversation1 = ''
with open(picked_case['ctpt'], 'r', encoding='utf-8') as file:
    conversation1 = file.read()

conversation2 = ''
with open(picked_case['solo'], 'r', encoding='utf-8') as file:
    file_content = file.read()

email1 = read_word('../data/physio/Gavin/Gavin Email.docx')
email2 = read_word('../data/physio/Margrate/Margaret Demo Email.docx')

progress_note = read_word(picked_case['progress_note'])

primary_agent = Agent(
    get_model(),
    system_prompt=f"""
    You are a Senior Physiotherapist who is very experienced in writing email to the care coordinator stakeholder based on conversation info and progress note.
    You will have one conversation with patient, then have solo conversation which helps to recall the key content.
    You will also have progress note.
    
    Here are some sample emails for you to learn how to write to the care coordinator stakeholder based on conversations and progress note:
        1, email sample 1: {email1} .
        2, email sample 2: {email2} .

    """
)

content = f"""
    Here are two conversations:
        Conversation 1, which is between a physiotherapist and a patient: {conversation1} .
        Conversation 2, which is physiotherapist solo: {conversation2} .
    Here is a progress note:
        {progress_note} .

    Please write a professional email to the care coordinator stakeholder with word friendly format.  
"""
# print(file_content)

result = primary_agent.run_sync(content)
email_result = result.data


# File paths
write_text_to_word(email_result,picked_case['output_file'])
print('done')

from __future__ import annotations
from contextlib import AsyncExitStack
from typing import Any, Dict, List
from dataclasses import dataclass
from dotenv import load_dotenv
from rich.markdown import Markdown
from rich.console import Console
from rich.live import Live
import asyncio
import os

from pydantic import BaseModel, Field
from pydantic_ai.providers.openai import OpenAIProvider
from pydantic_ai.models.openai import OpenAIModel
from pydantic_ai.mcp import MCPServerStdio
from pydantic_ai import Agent, RunContext

from docx import Document
from docx.shared import Inches
import docx
import re
import os

load_dotenv()


gavin_case = {
    'ctpt': '../data/physio/Gavin/gavinpt.txt',
    'solo': '../data/physio/Gavin/gavinsolo.txt',
    'progress_note' : r"../data/physio/Gavin/Progress_Note_Gavin_updated.docx",
    'output_file' : r"../data/physio/Gavin/Email_Gavin.docx",
    'template_file' : r"../data/physio/physiotherapy_email_template.docx"
}

margrate_case = {
    'ctpt': '../data/physio/Margrate/margaretpt.txt',
    'solo': '../data/physio/Margrate/margaretsolo.txt',
    'progress_note' : r"../data/physio/Margrate/Progress_Note_Margrate_updated.docx",
    'output_file' : r"../data/physio/Margrate/Email_Margrate.docx",
    'template_file' : r"../data/physio/physiotherapy_email_template.docx",
}

test_case = {
    'ctpt': '../data/physio/Test/testpt.txt',
    'solo': '../data/physio/Test/testsolo.txt',
    'progress_note' : r"../data/physio/Test/Progress_Note_Test_updated.docx",
    'output_file' : r"../data/physio/Test/Email_Test.docx",
    'template_file' : r"../data/physio/physiotherapy_email_template.docx",
}

# ========== Helper function to get model configuration ==========
def get_g_model():
    llm = 'o3'
    # print(llm)
    base_url = 'https://api.openai.com/v1'
    # print(base_url)
    api_key = os.getenv("OPENAI_API_KEY")
    # print(api_key)
    return OpenAIModel(llm, provider=OpenAIProvider(base_url=base_url, api_key=api_key))

def get_o_model():
    llm = 'google/gemini-2.5-pro'
    base_url = 'https://openrouter.ai/api/v1'
    api_key = os.getenv("OPEN_ROUTER_API_KEY")
    return OpenAIModel(llm, provider=OpenAIProvider(base_url=base_url, api_key=api_key))


def update_word_template(template_path, output_path, replacement_dict, image_replacements):
    doc = Document(template_path)
    pattern = r'\[(.*?)\]'

    def process_paragraphs(paragraphs):
        for paragraph in paragraphs:
            full_text = ''.join(run.text for run in paragraph.runs)

            if '[Scooter_Image]' in full_text:
                print("found image")


                # Clear all runs
                for run in paragraph.runs:
                    run.text = ''

                # Split around the image placeholder
                parts = full_text.split('[Scooter_Image]')
                
                # Rebuild the paragraph
                if parts[0]:
                    paragraph.add_run(parts[0])
                # Add the image
                paragraph.add_run().add_picture(image_replacements['Scooter_Image'], width=Inches(3))
                if len(parts) > 1:
                    paragraph.add_run(parts[1])
            else:
                for match in re.finditer(pattern, paragraph.text):
                    key = match.group(1)
                    if key in replacement_dict:
                        paragraph.text = paragraph.text.replace(f'[{key}]', replacement_dict[key])

    process_paragraphs(doc.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

    for section in doc.sections:
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer is not None:
                for p in footer.paragraphs:
                    print("Footer paragraph:", p.text)

                process_paragraphs(footer.paragraphs)
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            # for p in cell.paragraphs:
                            #     print("Cell paragraph:", p.text)
                            process_paragraphs(cell.paragraphs)

    # for section in doc.sections:
    #     footer = section.footer
    #     # for p in footer.paragraphs:
    #     #     print("Footer paragraph:", p.text)
        
    #     process_paragraphs(footer.paragraphs)

    #     # If your footer contains tables
    #     for table in footer.tables:
    #         for row in table.rows:
    #             for cell in row.cells:
    #                 # for p in cell.paragraphs:
    #                 #     print("Cell paragraph:", p.text)
    #                 process_paragraphs(cell.paragraphs)


    for section in doc.sections:
        header = section.header
        process_paragraphs(header.paragraphs)

        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    process_paragraphs(cell.paragraphs)



    doc.save(output_path)
    print(f"Document saved to {output_path}")

def write_text_to_word(text, filename="document.docx", title=None):
    """
    Write text to a Word document.
    
    Args:
        text (str): The text content to write to the document
        filename (str): Name of the output file (default: "document.docx")
        title (str, optional): Optional title for the document
    
    Returns:
        str: Path to the created document
    """
    # Create a new Document
    doc = Document()
    
    # Add title if provided
    if title:
        title_paragraph = doc.add_heading(title, 0)
    
    # Add the main text content
    # Split text by paragraphs (double newlines) for better formatting
    paragraphs = text.split('\n\n')
    
    for paragraph_text in paragraphs:
        if paragraph_text.strip():  # Skip empty paragraphs
            doc.add_paragraph(paragraph_text.strip())
    
    # Save the document
    doc.save(filename)
    
    return filename

def read_word(file_path):
    """
    read text from word
    """
    try:
        # Load the Word document
        doc = docx.Document(file_path)
        
        # Extract all text from the document
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        
        text = '\n'.join(full_text)
        
        return text
    
    except Exception as e:
        return f"Error processing file: {str(e)}"
    

    
    


picked_case = margrate_case
selected_model = get_g_model()





class EmailContent(BaseModel):
    first_name_patient: str = Field(description="Patient’s first name used in the email.")
    last_name_patient: str = Field(description="Patient’s last name used in the email.")
    hello_info: str = Field(description="Opening greeting line to the support coordinator (e.g., 'Good morning [Name]').")
    greeting_info: str = Field(description="Warm follow-up greeting line to create rapport (e.g., 'I hope you're having a good day!').")
    date_visit: str = Field(description="Date the patient was seen for their initial physiotherapy session.")
    main_issue: str = Field(description="Primary issue affecting the patient’s function and daily living. Use full sentences or paragraph format.")
    other_issue: str = Field(description="Secondary or less critical issues mentioned during assessment. Use full sentences or paragraph format.")
    Physiotherapist_suggestion: str = Field(description="Physiotherapist's clinical recommendations such as referrals to other allied health services (e.g., OT review).")
    outcome_fact_measures_vs_mean_values_then_major_risk_indication: str = Field(
        description="List of outcome measures taken (e.g., STS, TUGT, 10MWT), comparison with normative data, and interpretation regarding fall or functional risk. Use bullet points if needed for clarity."
    )
    goals_from_patient: str = Field(
        description="Client-identified goals for physiotherapy, including short- and long-term functional aims. Use bullet points or paragraphs for clarity."
    )
    current_status: str = Field(
        description="Patient’s current status in terms of endurance, tolerance to activity, cooperation, or barriers identified during the session. Use paragraph format."
    )
    recomendations: str = Field(
        description="Follow-up actions, frequency of physiotherapy, requests (e.g., medical history), and suggestions for further support or referrals. Use bullet points or clear sentence structure."
    )
    date_next_appointment: str = Field(
        description="Tentative date and time for the next physiotherapy session."
    )
    first_name_physiotherapist: str = Field(description="Physiotherapist’s first name.")
    last_name_physiotherapist: str = Field(description="Physiotherapist’s last name.")


    
conversation1 = ''
with open(picked_case['ctpt'], 'r', encoding='utf-8') as file:
    conversation1 = file.read()

conversation2 = ''
with open(picked_case['solo'], 'r', encoding='utf-8') as file:
    file_content = file.read()

email1 = read_word('../data/physio/Gavin/Gavin Email.docx')
email2 = read_word('../data/physio/Margrate/Margaret Demo Email.docx')

progress_note = read_word(picked_case['progress_note'])

primary_agent = Agent(
    selected_model,
    system_prompt=f"""
    You are a Senior Physiotherapist who is very experienced in writing email to the care coordinator stakeholder based on progress note and conversation info and progress note.
    You will have one conversation with patient, and have solo contents which helps to recall the key contents.
    You will have the progress note.
    
    Here are some sample emails for you to learn how to write to the care coordinator stakeholder based on conversations and progress note:
        1, email sample 1: {email1} .
        2, email sample 2: {email2} .
    End of sample emails.

    You need to try to use terminology to be professional and simplify description wherever it is possible, here are some terminologies for you to know:
    1, Use FAEO to stand for "Feet Apart, Eyes Open". 
    2, Use FAEC to stand for "Feet Apart, Eyes Closed". 
    3, Use FTEO to stand for "Feet Together, Eyes Open". 
    4, Use FTEC to stand for "Feet Together, Eyes Closed".
    5, Use UL to stand for "Upper Limbs". 
    6, Use TF to stand for "Task/Function".
    7, Use 5x STS to stand for "Five Times Sit-to-Stand Test". 
    8, Use 4WW to stand for "4-Wheeled Walker".
    9, Use WC to stand for "Wheelchair".
    10, Use R/V 1/52 to mean "Review in 1 week".
    11, Use A+O to stand for "Alert and Oriented".
    12, Use TPP to stand for "Time, Place, Person", indicating awareness of these three elements.
    13, Use I/M to stand for "Intermittent", meaning it occurs from time to time.
    14, Use Dx to stand for "Diagnosis".
    15, Use Hx to stand for "History".
    16, Use a/a to mean "as above".
    17, Use 'Formal Dx ~6/12 ago' to mean a formal diagnosis was made about 6 months ago.
    18, Use Cx to stand for "Cervical", referring to the neck region of the spine.
    19, Use Tx to stand for "Thoracic", referring to the mid-back region of the spine.
    20, Use Lx to stand for "Lumbar", referring to the lower back region of the spine.
    21, Use ROM to stand for "Range of Motion", which measures how far a joint can move.
    22, Use A = 80/80 to mean the joint can actively move through the full expected range of 80 degrees.
    23, Use P = 80/80 to mean the joint can passively be moved through the full expected range of 80 degrees.
    24, Use A = ½ to mean the joint can only move through half of the expected active range.
    25, Use 'Modified 30s: 4x' to mean a modified version of the 30-second sit-to-stand test where the patient was able to perform the movement 4 times within 30 seconds.
    26, Use 0.47 m/s to describe a walking speed of 0.47 meters per second.
    27, Use Ax to stand for "Assessment".
    28, Use BBS to stand for "Berg Balance Scale", a tool used to assess balance.
    29, Use PMS to stand for "Physical Mobility Scale".
    30, Use Cont. to mean "Continue".
    31, Use PMHx to stand for "Past Medical History".
    32, Use SC to stand for either "Specialist Consultant" or "Service Coordinator", depending on the context.
    33, Use PMD to stand for "Personal Mobility Device".
    34, Use ERC to stand for "Equipment Resource Center".
    35, Use PMS to stand for "Performance Measurement System" or "Patient Management System", depending on context.
    36, Use CV to stand for "Cardiovascular".
    37, Use OT to stand for "Occupational Therapy".
    38, Use RV to stand for "Review".
    39, Use PT to stand for "Physiotherapy" (or Physical Therapy).
    40, Use COPD to stand for "Chronic Obstructive Pulmonary Disease", a chronic lung condition that causes long-term breathing problems.
    41, Use ADLs to stand for "Activities of Daily Living", such as eating, bathing, dressing, and moving around.
    42, Use SOB to stand for "Shortness of Breath", a symptom often linked with breathing difficulties.
    43, Use OM to stand for "Outcome Measures", which are tools used to assess patient progress and effectiveness of treatment.

    by looking at the progress note and conversations which will be provide, you need to extract the information and fill in the following fields:

    1, first_name_patient: "Patient’s first name used in the email."
    2, last_name_patient: "Patient’s last name used in the email."
    3, hello_info: "Opening greeting line to the support coordinator (e.g., 'Good morning [Name]')."
    4, greeting_info: "Warm follow-up greeting line to create rapport (e.g., 'I hope you're having a good day!')."
    5, date_visit: "Date the patient was seen for their initial physiotherapy session."
    6, main_issue: "Primary issue affecting the patient’s function and daily living. Use full sentences or paragraph format."
    7, other_issue: "Secondary or less critical issues mentioned during assessment. Use full sentences or paragraph format."
    8, Physiotherapist_suggestion: "Physiotherapist's clinical recommendations such as referrals to other allied health services (e.g., OT review)."
    9, outcome_fact_measures_vs_mean_values_then_major_risk_indication: "List of outcome measures taken (e.g., STS, TUGT, 10MWT), comparison with normative data (you can search web for info), and interpretation regarding fall or functional risk. Use bullet points if needed for clarity."
    10, goals_from_patient: "Client-identified goals for physiotherapy, including short- and long-term functional aims. Use bullet points or paragraphs for clarity."
    11, current_status: "Patient’s current status in terms of endurance, tolerance to activity, cooperation, or barriers identified during the session. Use paragraph format."
    12, recomendations: "Follow-up actions, frequency of physiotherapy, requests (e.g., medical history), and suggestions for further support or referrals. Use bullet points or clear sentence structure."
    13, date_next_appointment: "Tentative date and time for the next physiotherapy session."
    14, first_name_physiotherapist: "Physiotherapist’s first name."
    15, last_name_physiotherapist: "Physiotherapist’s last name."

    """,
    result_type=EmailContent,
)

content = f"""
    Here is one conversation and solo content:
        Conversation 1, which is between a physiotherapist and a patient: {conversation1} .
        Conversation 2, which is physiotherapist solo: {conversation2} .
    Here is a progress note:
        {progress_note} .

    Please write a professional email to the care coordinator stakeholder with word friendly format.  
"""
# print(file_content)

result = primary_agent.run_sync(content)
email_result = result.data

# File paths
print(email_result.first_name_patient)

replacements = {
    'first_name_patient': email_result.first_name_patient,
    'last_name_patient': email_result.last_name_patient,
    'hello_info': email_result.hello_info,
    'greeting_info': email_result.greeting_info,
    'date_visit': email_result.date_visit,
    'main_issue': email_result.main_issue,
    'other_issue': email_result.other_issue,
    'Physiotherapist_suggestion': email_result.Physiotherapist_suggestion,
    'outcome_fact_measures_vs_mean_values_then_major_risk_indication': email_result.outcome_fact_measures_vs_mean_values_then_major_risk_indication,
    'goals_from_patient': email_result.goals_from_patient,
    'current_status': email_result.current_status,
    'recomendations': email_result.recomendations,
    'date_next_appointment': email_result.date_next_appointment,
    'first_name_physiotherapist': email_result.first_name_physiotherapist,
    'last_name_physiotherapist': email_result.last_name_physiotherapist
}


# Provide image path for [Scooter_Image]
image_replacements = {
    'Signature_Image': '../img/scooter.png'  # Replace with your actual image path
}

# File paths
input_path = picked_case['template_file']
output_path = picked_case['output_file']

update_word_template(input_path, output_path, replacements, image_replacements)



import docx
import re
from collections import Counter

def count_tokens_basic(file_path):
    """
    Basic word tokenization - splits on whitespace and punctuation
    """
    try:
        # Load the Word document
        doc = docx.Document(file_path)
        
        # Extract all text from the document
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        
        text = '\n'.join(full_text)
        
        # Basic tokenization - split on whitespace and remove punctuation
        tokens = re.findall(r'\b\w+\b', text.lower())
        
        return {
            'total_tokens': len(tokens),
            'unique_tokens': len(set(tokens)),
            'tokens': tokens,
            'token_frequency': Counter(tokens)
        }
    
    except Exception as e:
        return f"Error processing file: {str(e)}"

def count_tokens_advanced(file_path):
    """
    More advanced tokenization using NLTK (requires: pip install nltk)
    """
    try:
        import nltk
        from nltk.tokenize import word_tokenize
        from nltk.corpus import stopwords
        
        # Download required NLTK data (run once)
        # nltk.download('punkt')
        # nltk.download('stopwords')
        
        # Load the Word document
        doc = docx.Document(file_path)
        
        # Extract all text
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        
        text = ' '.join(full_text)
        
        # Tokenize using NLTK
        tokens = word_tokenize(text.lower())
        
        # Remove punctuation and non-alphabetic tokens
        tokens = [token for token in tokens if token.isalpha()]
        
        # Optional: Remove stopwords
        stop_words = set(stopwords.words('english'))
        filtered_tokens = [token for token in tokens if token not in stop_words]
        
        return {
            'total_tokens': len(tokens),
            'total_tokens_no_stopwords': len(filtered_tokens),
            'unique_tokens': len(set(tokens)),
            'unique_tokens_no_stopwords': len(set(filtered_tokens)),
            'tokens': tokens,
            'filtered_tokens': filtered_tokens,
            'token_frequency': Counter(tokens),
            'filtered_frequency': Counter(filtered_tokens)
        }
    
    except ImportError:
        return "NLTK not installed. Run: pip install nltk"
    except Exception as e:
        return f"Error processing file: {str(e)}"

def count_tokens_sentences(file_path):
    """
    Count sentences and words per sentence
    """
    try:
        # Load the Word document
        doc = docx.Document(file_path)
        
        # Extract all text
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        
        text = ' '.join(full_text)
        
        # Split into sentences
        sentences = re.split(r'[.!?]+', text)
        sentences = [s.strip() for s in sentences if s.strip()]
        
        sentence_tokens = []
        for sentence in sentences:
            tokens = re.findall(r'\b\w+\b', sentence.lower())
            sentence_tokens.append({
                'sentence': sentence,
                'token_count': len(tokens),
                'tokens': tokens
            })
        
        return {
            'total_sentences': len(sentences),
            'sentence_analysis': sentence_tokens,
            'avg_tokens_per_sentence': sum(s['token_count'] for s in sentence_tokens) / len(sentence_tokens) if sentence_tokens else 0
        }
    
    except Exception as e:
        return f"Error processing file: {str(e)}"

# Usage examples
if __name__ == "__main__":
    file_path = "./data/Gavin Progress Note.docx"  # Replace with your file path
    
    print("=== Basic Token Count ===")
    basic_result = count_tokens_basic(file_path)
    if isinstance(basic_result, dict):
        print(f"Total tokens: {basic_result['total_tokens']}")
        print(f"Unique tokens: {basic_result['unique_tokens']}")
        print(f"Most common tokens: {basic_result['token_frequency'].most_common(10)}")
    else:
        print(basic_result)
    
    print("\n=== Advanced Token Count (with NLTK) ===")
    advanced_result = count_tokens_advanced(file_path)
    if isinstance(advanced_result, dict):
        print(f"Total tokens: {advanced_result['total_tokens']}")
        print(f"Unique tokens: {advanced_result['unique_tokens']}")
        print(f"Tokens without stopwords: {advanced_result['total_tokens_no_stopwords']}")
        print(f"Most common tokens: {advanced_result['token_frequency'].most_common(10)}")
    else:
        print(advanced_result)
    
    print("\n=== Sentence Analysis ===")
    sentence_result = count_tokens_sentences(file_path)
    if isinstance(sentence_result, dict):
        print(f"Total sentences: {sentence_result['total_sentences']}")
        print(f"Average tokens per sentence: {sentence_result['avg_tokens_per_sentence']:.2f}")
    else:
        print(sentence_result)
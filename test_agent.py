from __future__ import annotations
from contextlib import AsyncExitStack
from typing import Any, Dict, List
from dataclasses import dataclass
from dotenv import load_dotenv
from rich.markdown import Markdown
from rich.console import Console
from rich.live import Live
import asyncio
import os

from pydantic import BaseModel, Field
from pydantic_ai.providers.openai import OpenAIProvider
from pydantic_ai.models.openai import OpenAIModel
from pydantic_ai.mcp import MCPServerStdio
from pydantic_ai import Agent, RunContext

from docx import Document
from docx.shared import Inches
import re
import os

load_dotenv()


def update_word_template(template_path, output_path, replacement_dict, image_replacements):
    doc = Document(template_path)
    pattern = r'\[(.*?)\]'

    def process_paragraphs(paragraphs):
        for paragraph in paragraphs:
            full_text = ''.join(run.text for run in paragraph.runs)

            if '[Scooter_Image]' in full_text:
                print("found image")


                # Clear all runs
                for run in paragraph.runs:
                    run.text = ''

                # Split around the image placeholder
                parts = full_text.split('[Scooter_Image]')
                
                # Rebuild the paragraph
                if parts[0]:
                    paragraph.add_run(parts[0])
                # Add the image
                paragraph.add_run().add_picture(image_replacements['Scooter_Image'], width=Inches(3))
                if len(parts) > 1:
                    paragraph.add_run(parts[1])
            else:
                for match in re.finditer(pattern, paragraph.text):
                    key = match.group(1)
                    if key in replacement_dict:
                        paragraph.text = paragraph.text.replace(f'[{key}]', replacement_dict[key])

    process_paragraphs(doc.paragraphs)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)

    for section in doc.sections:
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer is not None:
                for p in footer.paragraphs:
                    print("Footer paragraph:", p.text)

                process_paragraphs(footer.paragraphs)
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            # for p in cell.paragraphs:
                            #     print("Cell paragraph:", p.text)
                            process_paragraphs(cell.paragraphs)

    # for section in doc.sections:
    #     footer = section.footer
    #     # for p in footer.paragraphs:
    #     #     print("Footer paragraph:", p.text)
        
    #     process_paragraphs(footer.paragraphs)

    #     # If your footer contains tables
    #     for table in footer.tables:
    #         for row in table.rows:
    #             for cell in row.cells:
    #                 # for p in cell.paragraphs:
    #                 #     print("Cell paragraph:", p.text)
    #                 process_paragraphs(cell.paragraphs)


    for section in doc.sections:
        header = section.header
        process_paragraphs(header.paragraphs)

        for table in header.tables:
            for row in table.rows:
                for cell in row.cells:
                    process_paragraphs(cell.paragraphs)



    doc.save(output_path)
    print(f"Document saved to {output_path}")


class Assessment(BaseModel):
    client: str = Field(description='patient or client name', )
    DOB: str = Field(description='date of birth')
    Address: str = Field(description='patient or client address')
    DOA: str = Field(description='date of assessment')
    DOR: str = Field(description='date of referral')
    Assessor: str = Field(description='assessor or therapist name')
    Qualification: str = Field(description='assessor or therapist qualification')
    Referral_Information: str = Field(description='summary of referral information as to why the patient was referred, planning')
    Background: str = Field(description='patient or client background, current situation, and any other relevant information, then use bullet points to list relevant information')
    Diagnosis_and_Conditions: str = Field(description='patient or client diagnosis and conditions, use bullet points to list relevant information')
    Assessment: str = Field(description='results of the assessment, patient or client needs, occupational therapy efforts, observations, advices, recommendations')
    Scooter_Name: str = Field(description='name of the scooter')
    Scooter_Description: str = Field(description='search key features of the scooter, then using bullet points to list the key features')
    Summary_and_Recommendation: str = Field(description='summary and recommendation')

# ========== Helper function to get model configuration ==========
def get_model():
    llm = 'gpt-4o'
    print(llm)
    base_url = 'https://api.openai.com/v1'
    print(base_url)
    api_key = ''
    print(api_key)
    return OpenAIModel(llm, provider=OpenAIProvider(base_url=base_url, api_key=api_key))


primary_agent = Agent(
    get_model(),
    system_prompt="""
    you are a senior occupational therapist who is very experienced in Occupational Therapy Assessment with patient and writing Equipment Recommendation Report:
    by looking at the conversations, you need to extract the information and fill in the following fields:

    1, client: 'patient or client name'
    2, DOB: 'date of birth'
    3, Address: 'patient or client address'
    4, DOA: 'date of assessment'
    5, DOR: 'date of referral'
    6, Assessor: 'assessor or therapist name'
    7, Qualification: 'assessor or therapist qualification'
    8, Referral_Information: 'summary of referral information as to why the patient was referred, planning'
    9, Background: 'patient or client background, current situation, and any other relevant information, then use bullet points to list relevant information'
    10, Diagnosis_and_Conditions: 'patient or client diagnosis and conditions, use bullet points to list relevant information'
    11, Assessment: 'results of the assessment, patient or client needs, occupational therapy efforts, observations, advices, recommendations'
    12, Scooter_Name: 'name of the scooter'
    13, Scooter_Description: 'search key features of the scooter, then using bullet points to list the key features'
    14, Summary_and_Recommendation: 'summary and recommendation'
    """,
    result_type=Assessment,
)

with open(r"./data/001.txt", 'r', encoding='utf-8') as file:
    file_content = file.read()

# print(file_content)

result = primary_agent.run_sync(file_content)
assess_result = result.data

print(assess_result.client)


# Dictionary for text replacements
replacements = {
    'client': assess_result.client,
    'DOB': assess_result.DOB,
    'Address': assess_result.Address,
    'DOA': assess_result.DOA,
    'DOR': assess_result.DOR,
    'Assessor': assess_result.Assessor,
    'Qualification': assess_result.Qualification,
    'Referral_Information': assess_result.Referral_Information,
    'Background': assess_result.Background,
    'Diagnosis_and_Conditions': assess_result.Diagnosis_and_Conditions,
    'Assessment': assess_result.Assessment,
    'Scooter_Name': 'Rothcare Boston Scooter',
    'Scooter_Description': assess_result.Scooter_Description,
    'Summary_and_Recommendation': assess_result.Summary_and_Recommendation,
    'Signature': assess_result.Assessor,
    'Assessor_info': 'p 08 7095 7904 | e jane.smith@regentherapy.com.au\nPO Box 3337, NORWOOD, SA 5067 | www.regentherapy.com.au',
    'Therapy Info': 'ReGen Therapy\nABN: 74 637 611 422'
}

# Provide image path for [Scooter_Image]
image_replacements = {
    'Scooter_Image': './img/scooter.png'  # Replace with your actual image path
}

# File paths
input_path = r"C:\Projects\prag\data\001_template.docx"
output_path = r"C:\Projects\prag\data\001_updated.docx"

update_word_template(input_path, output_path, replacements, image_replacements)